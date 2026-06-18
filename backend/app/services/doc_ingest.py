"""文件智慧匯入:任意文件(PDF / Word / Excel / CSV / 文字)→ 純文字 → Claude 抽取
結構化訂單 → 交回現有匯入流程(建單 + 地理編碼)。

設計取捨:
- **文字抽取**刻意用輕量原生套件(pypdf / python-docx / openpyxl / xlrd / 純解碼),
  不引入 MarkItDown 全家桶(其 magika + onnxruntime 過重),符合本專案「精簡」策略;
  抽取層為單一函式,日後要換 MarkItDown 只需替換 `extract_text`。
- ⚠️ **個資**:本流程會把文件原文(可能含乘客姓名/電話/地址)送往 Claude API 抽取。
  與「送 Claude 一律去識別化」原則有張力——文件抽取本質需要原文。正式上線處理真實 PII 前,
  應改用本地/地端模型抽取(見 BACKLOG「正式部署:個資合規」)。此處不記錄文件內容到日誌。
"""
from __future__ import annotations

import io
import json
import re
from datetime import date

from app.services import ai_dispatch

SUPPORTED_EXT = {"pdf", "docx", "xlsx", "xlsm", "xls", "csv", "txt", "md"}

# 抽取輸出允許的訂單欄位(對齊 OrderCreate)
_ALLOWED = {
    "service_date", "pickup_time", "pickup_window_min", "passenger_name",
    "passenger_phone", "pickup_address", "dropoff_address", "pax",
    "vehicle_type", "need_wheelchair", "allow_pool", "note",
}
_REQUIRED = ("service_date", "pickup_time", "pickup_address", "dropoff_address")


def _ext(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def extract_text(filename: str, content: bytes) -> str:
    """把上傳文件轉成純文字。不支援的副檔名拋 ValueError。"""
    ext = _ext(filename)
    if ext == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    if ext == "docx":
        import docx
        d = docx.Document(io.BytesIO(content))
        parts = [p.text for p in d.paragraphs]
        for tbl in d.tables:
            for row in tbl.rows:
                parts.append("\t".join(c.text for c in row.cells))
        return "\n".join(parts)
    if ext in ("xlsx", "xlsm"):
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        out = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                out.append("\t".join("" if c is None else str(c) for c in row))
        return "\n".join(out)
    if ext == "xls":
        import xlrd
        book = xlrd.open_workbook(file_contents=content)
        out = []
        for sh in book.sheets():
            for r in range(sh.nrows):
                out.append("\t".join(str(sh.cell_value(r, c)) for c in range(sh.ncols)))
        return "\n".join(out)
    if ext in ("csv", "txt", "md"):
        for enc in ("utf-8-sig", "utf-8", "big5", "cp950"):
            try:
                return content.decode(enc)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="replace")
    raise ValueError(f"不支援的檔案格式:.{ext}(支援:{', '.join(sorted(SUPPORTED_EXT))})")


_SYSTEM = (
    "你是長照接送派遣的資料抽取助手。使用者會貼上一份文件的純文字內容,"
    "其中包含一筆或多筆接送訂單。請只輸出 JSON 陣列(不要任何說明文字、不要 markdown 圍欄)。"
)


def _prompt(text: str, default_date: str | None) -> str:
    dd = f"若文件未標明日期,服務日期一律用 {default_date}。" if default_date else ""
    return (
        "請從以下文件抽取所有接送訂單,輸出 JSON 陣列,每筆物件欄位如下(未知用 null):\n"
        '- service_date: "YYYY-MM-DD"\n'
        '- pickup_time: "YYYY-MM-DDTHH:MM:00"(若只有時間就用該服務日期組合)\n'
        "- passenger_name: 字串\n- passenger_phone: 字串\n"
        "- pickup_address: 字串(上車地址,盡量完整含縣市)\n"
        "- dropoff_address: 字串(下車地址)\n"
        "- pax: 整數(人數,預設 1)\n"
        '- vehicle_type: "welfare"(福祉車/輪椅) 或 "normal"(一般車)\n'
        "- need_wheelchair: true/false\n- allow_pool: true/false(可否共乘,未提及給 false)\n"
        "- note: 字串(其他備註)\n"
        f"{dd}\n只輸出 JSON 陣列。文件內容如下:\n---\n{text}\n---"
    )


def _strip_json(raw: str) -> str:
    """去除可能的 markdown 圍欄,取出 JSON 陣列主體。"""
    s = raw.strip()
    s = re.sub(r"^```(?:json)?\s*|\s*```$", "", s, flags=re.IGNORECASE).strip()
    i, j = s.find("["), s.rfind("]")
    return s[i:j + 1] if i != -1 and j != -1 and j > i else s


def _coerce(rec: dict, default_date: str | None) -> dict:
    """把單筆抽取結果正規化成 OrderCreate 可接受的 payload。"""
    o = {k: rec.get(k) for k in _ALLOWED if rec.get(k) not in (None, "")}
    if "service_date" not in o and default_date:
        o["service_date"] = default_date
    pt = o.get("pickup_time")
    if pt is not None:
        pt = str(pt).strip()
        # 只有 HH:MM → 與服務日期組合
        if re.fullmatch(r"\d{1,2}:\d{2}", pt) and o.get("service_date"):
            hh, mm = pt.split(":")
            o["pickup_time"] = f"{o['service_date']}T{int(hh):02d}:{mm}:00"
        else:
            o["pickup_time"] = pt
    vt = str(o.get("vehicle_type", "")).strip().lower()
    if vt:
        o["vehicle_type"] = "welfare" if vt in (
            "welfare", "福祉", "福祉車", "wheelchair", "輪椅") else "normal"
    for b in ("need_wheelchair", "allow_pool"):
        if b in o and not isinstance(o[b], bool):
            o[b] = str(o[b]).strip().lower() in ("1", "true", "yes", "是", "y", "v")
    if "pax" in o:
        try:
            o["pax"] = int(float(o["pax"]))
        except (ValueError, TypeError):
            o.pop("pax")
    return o


def extract_orders(text: str, default_date: str | None = None) -> tuple[list[dict], list[dict], str]:
    """文字 → Claude 抽取 → (payloads, errors, raw)。payloads 為可建單的 dict。"""
    if not text.strip():
        return [], [{"error": "文件無可讀文字"}], ""
    raw = ai_dispatch._call_claude(
        _prompt(text, default_date), system=_SYSTEM, max_tokens=4096, timeout=90)
    try:
        data = json.loads(_strip_json(raw))
    except json.JSONDecodeError as e:
        return [], [{"error": f"AI 回傳非有效 JSON:{e}"}], raw
    if isinstance(data, dict):
        data = [data]
    if not isinstance(data, list):
        return [], [{"error": "AI 回傳格式非陣列"}], raw

    payloads, errors = [], []
    for i, rec in enumerate(data):
        if not isinstance(rec, dict):
            errors.append({"row": f"#{i + 1}", "error": "非物件"})
            continue
        o = _coerce(rec, default_date)
        missing = [f for f in _REQUIRED if not o.get(f)]
        if missing:
            errors.append({"row": f"#{i + 1}", "error": f"缺必要欄位:{', '.join(missing)}",
                           "data": {k: rec.get(k) for k in _REQUIRED}})
            continue
        payloads.append(o)
    return payloads, errors, raw
